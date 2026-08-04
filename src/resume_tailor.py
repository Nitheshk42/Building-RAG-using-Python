import re
import difflib
import io
import os
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document as DocxDocument


def _get_llm(temperature=0.4, max_tokens=3000):
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY missing! Check .env file exists at project root")
    return ChatGroq(
        model="llama-3.3-70b-versatile",
        temperature=temperature,
        max_tokens=max_tokens,
        api_key=api_key
    )


def analyze_resume_for_jd(resume_text, jd_text):
    """Finds the first two project sections in the resume and, against the JD, produces:
    - SUGGESTIONS: new candidate bullet points to ADD (existing text is never rewritten or
      touched). Each suggestion is evaluated individually - if a quantitative metric would
      genuinely strengthen that specific point, an illustrative estimated range is woven into
      its text; if not, it's left without one. Nothing is forced.
    - MISSING_SKILLS: JD-required tools/tech not evidenced in the resume, each with an optional
      draft bullet if there's an honest, plausible tie-in to real project context.
    Nothing is ever auto-applied - the original text is only changed once the user checks a
    suggestion on, and only that approved text gets added."""
    llm = _get_llm()
    template = """You are a resume coach. Below is a candidate's RESUME and a JOB DESCRIPTION.

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

Find the FIRST TWO distinct project/experience entries in the resume (in the order they
appear). For EACH of the two, produce:

1. ORIGINAL: the verbatim text of that project section, copied exactly from the resume, so it
   can be matched back against the original document. Do NOT reword or restructure this — it
   must be an exact copy.

2. SUGGESTIONS: 3-5 NEW candidate bullet points that could be ADDED to this project (never
   modifications to existing lines — existing text is never touched). Each must be a ready-to-
   insert, first-person, action-verb resume bullet, plausible from context and relevant to the
   JD, never claiming a tool/platform the candidate has no evidence of using.

   For EACH suggestion, decide individually whether a quantitative metric would genuinely
   strengthen THAT specific point (e.g. a bullet about optimization or scale usually benefits
   from a number; a bullet about collaboration or process usually doesn't). If yes, weave a
   plausible ILLUSTRATIVE ESTIMATED RANGE directly into that bullet's sentence (e.g. "...cutting
   processing time by roughly 25-30%"). If a metric doesn't naturally fit that particular point,
   leave it out entirely — do not force a number into every bullet. These ranges are estimates
   for the candidate to review and replace with their real figure, never confirmed facts.

3. MISSING_SKILLS: for each tool/language/platform the JOB DESCRIPTION asks for that is NOT
   evidenced anywhere in this project or the resume, output ONE LINE in this exact pipe format:
   SKILL :: NOTE :: DRAFT
   - SKILL: the tool/tech name
   - NOTE: one sentence flagging this as a real gap, not something to fabricate
   - DRAFT: a draft bullet blending this skill into the REAL project context ONLY IF plausible
     (phrase speculatively, e.g. "Could describe: ..."). If no plausible tie-in exists, DRAFT
     should say "No plausible tie-in — do not add this."
   If no missing skills for this project, output: None :: None :: None

Respond in EXACTLY this format, nothing else:

PROJECT1_ORIGINAL:
<verbatim text>
PROJECT1_SUGGESTIONS:
- <suggestion, with an embedded metric range only if it genuinely fits>
PROJECT1_MISSING_SKILLS:
<SKILL :: NOTE :: DRAFT line(s), or "None :: None :: None">
===
PROJECT2_ORIGINAL:
<verbatim text>
PROJECT2_SUGGESTIONS:
- <suggestion, with an embedded metric range only if it genuinely fits>
PROJECT2_MISSING_SKILLS:
<SKILL :: NOTE :: DRAFT line(s), or "None :: None :: None">

Begin now:"""
    prompt = ChatPromptTemplate.from_template(template)
    chain = prompt | llm | StrOutputParser()
    result = chain.invoke({"resume_text": resume_text, "jd_text": jd_text})
    return _parse_projects(result)


_METRIC_PATTERN = re.compile(r"\d+\s*[-–]\s*\d+\s*%|\d+%|\d[\d,]*\s*(records|requests|users|gb|tb|ms|/day|/sec)", re.IGNORECASE)


def _bullets(text):
    return [line.strip("- ").strip() for line in text.strip().splitlines() if line.strip().startswith("-")]


def _parse_missing_skills(text):
    skills = []
    for line in text.strip().splitlines():
        line = line.strip().lstrip("-").strip()
        if not line or "::" not in line:
            continue
        parts = [p.strip() for p in line.split("::")]
        if len(parts) != 3:
            continue
        skill, note, draft = parts
        if skill.lower() == "none":
            continue
        skills.append({"skill": skill, "note": note, "draft": draft})
    return skills


def _parse_projects(raw):
    blocks = raw.split("===")
    projects = []
    pattern = re.compile(
        r"PROJECT\d+_ORIGINAL:\s*(?P<original>.*?)\s*"
        r"PROJECT\d+_SUGGESTIONS:\s*(?P<suggestions>.*?)\s*"
        r"PROJECT\d+_MISSING_SKILLS:\s*(?P<missing>.*)",
        re.DOTALL
    )
    for block in blocks:
        match = pattern.search(block)
        if not match:
            continue
        suggestions = _bullets(match.group("suggestions"))
        projects.append({
            "original": match.group("original").strip(),
            "suggestions": suggestions,
            "has_any_metric": any(_METRIC_PATTERN.search(s) for s in suggestions),
            "missing_skills": _parse_missing_skills(match.group("missing")),
        })
    return projects


def render_side_by_side_diff(original, edited):
    """GitHub PR-style split diff: returns (left_html, right_html). Left is the original with
    removed words struck through in red; right is the edited version with added words
    highlighted in green. Unchanged words render plainly on both sides."""
    orig_words = original.split()
    edit_words = edited.split()
    sm = difflib.SequenceMatcher(None, orig_words, edit_words)

    left_parts, right_parts = [], []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            left_parts.append(" ".join(orig_words[i1:i2]))
            right_parts.append(" ".join(edit_words[j1:j2]))
        elif tag == "delete":
            left_parts.append(
                f'<span style="background:#4a1414;color:#ff8a8a;text-decoration:line-through;padding:1px 3px;border-radius:3px;">{" ".join(orig_words[i1:i2])}</span>'
            )
        elif tag == "insert":
            right_parts.append(
                f'<span style="background:#143d1d;color:#8affa0;padding:1px 3px;border-radius:3px;">{" ".join(edit_words[j1:j2])}</span>'
            )
        elif tag == "replace":
            left_parts.append(
                f'<span style="background:#4a1414;color:#ff8a8a;text-decoration:line-through;padding:1px 3px;border-radius:3px;">{" ".join(orig_words[i1:i2])}</span>'
            )
            right_parts.append(
                f'<span style="background:#143d1d;color:#8affa0;padding:1px 3px;border-radius:3px;">{" ".join(edit_words[j1:j2])}</span>'
            )

    return " ".join(left_parts), " ".join(right_parts)


def build_tailored_docx(full_resume_text, replacements):
    """replacements: list of (original_snippet, final_text) tuples to substitute into the
    full resume text. Falls back to appending if an exact match isn't found (PDF text
    extraction can introduce whitespace differences)."""
    final_text = full_resume_text
    unmatched = []
    for original_snippet, final_snippet in replacements:
        if original_snippet in final_text:
            final_text = final_text.replace(original_snippet, final_snippet, 1)
        else:
            unmatched.append(final_snippet)

    doc = DocxDocument()
    for line in final_text.split("\n"):
        doc.add_paragraph(line)

    if unmatched:
        doc.add_paragraph("")
        doc.add_paragraph("--- Tailored additions (could not auto-place in original text) ---")
        for snippet in unmatched:
            doc.add_paragraph(snippet)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
